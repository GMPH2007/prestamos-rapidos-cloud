"""
=============================================================================
GENERADOR DE DOCUMENTACIÓN OFICIAL EN FORMATO MICROSOFT WORD (.DOCX)
=============================================================================
Este script compila toda la documentación técnica, manual de funcionamiento,
arquitectura cloud y entregables en un archivo Word (.docx) formal.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    """Aplica color de fondo a una celda de tabla en Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Ajusta márgenes internos de una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_word_document():
    doc = docx.Document()

    # Configuración de márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilos de Fuente
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(30, 41, 59)

    # =========================================================================
    # PORTADA OFICIAL
    # =========================================================================
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists('logo.png'):
        p_logo.add_run().add_picture('logo.png', width=Inches(1.6))

    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_univ = p_univ.add_run("INGENIERÍA DE SISTEMAS E INFORMÁTICA\nGESTIÓN DE PROYECTOS DE TI Y ARQUITECTURA CLOUD")
    run_univ.font.size = Pt(13)
    run_univ.font.bold = True
    run_univ.font.color.rgb = RGBColor(37, 99, 235)

    doc.add_paragraph("\n")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("TRABAJO ENCARGADO:\nPLATAFORMA WEB FINTECH DE PRÉSTAMOS RÁPIDOS Y PLAN DE MIGRACIÓN DE BASE DE DATOS A LA NUBE (AWS RDS)")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Manual de Funcionamiento, Arquitectura de Infraestructura, Matriz RACI y Ciclo de Vida del Proyecto")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("\n\n")

    # Tabla de Datos de Entrega
    table_meta = doc.add_table(rows=4, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Proyecto / Sistema:", "PrestaFast Cloud (Sistema Transaccional Web Móvil)"),
        ("Enlace Web en Vivo:", "https://gmph2007.github.io/prestamos-rapidos-cloud/"),
        ("Repositorio GitHub:", "https://github.com/GMPH2007/prestamos-rapidos-cloud"),
        ("Año Académico:", "2026")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = table_meta.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        r_lbl = cell_lbl.paragraphs[0].add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(30, 41, 59)
        
        r_val = cell_val.paragraphs[0].add_run(val)
        if "http" in val:
            r_val.font.color.rgb = RGBColor(37, 99, 235)
            r_val.font.underline = True
        
        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "F8FAFC")
        set_cell_margins(cell_lbl, 80, 80, 120, 120)
        set_cell_margins(cell_val, 80, 80, 120, 120)

    doc.add_page_break()

    # =========================================================================
    # SECCIÓN 1: RESUMEN EJECUTIVO Y OBJETIVOS
    # =========================================================================
    h1 = doc.add_heading("1. Resumen Ejecutivo del Proyecto", level=1)
    h1.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "El presente proyecto desarrolla una solución integral de tecnología financiera (Fintech) "
        "denominada PrestaFast Cloud, orientada a la colocación y desembolso digital de créditos de consumo "
        "en menos de 60 segundos. Como componente medular de infraestructura, se ejecuta el plan de "
        "migración de la base de datos relacional local (On-Premise) hacia una arquitectura gestionada en la nube "
        "con Amazon RDS PostgreSQL bajo un esquema de Alta Disponibilidad (Multi-AZ)."
    )

    doc.add_paragraph(
        "La solución elimina los cuellos de botella de hardware físico, reduce a cero el riesgo por cortes "
        "de energía o fallas de disco local y garantiza cumplimiento estricto de ciberseguridad con cifrado "
        "KMS AES-256 en reposo y TLS 1.3 en tránsito."
    )

    # =========================================================================
    # SECCIÓN 2: CLASIFICACIÓN DE PROYECTOS DE TI
    # =========================================================================
    h2 = doc.add_heading("2. Clasificación de los Tipos de Proyectos de TI Involucrados", level=1)
    h2.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "De acuerdo con los estándares internacionales del PMI (Project Management Institute) y TOGAF, "
        "esta iniciativa tecnológica abarca cuatro (4) tipologías complementarias de proyectos de TI:"
    )

    p_t1 = doc.add_paragraph()
    p_t1.add_run("1. Proyecto de Desarrollo de Software / Web App: ").font.bold = True
    p_t1.add_run("Construcción de la plataforma transaccional cliente/servidor responsiva con arquitectura Mobile-First, simulador interactivo de amortización francesa, motor de scoring y comprobantes digitales.")

    p_t2 = doc.add_paragraph()
    p_t2.add_run("2. Proyecto de Infraestructura y Migración Cloud: ").font.bold = True
    p_t2.add_run("Aprovisionamiento de la Red Privada Virtual (AWS VPC), balanceadores de carga elásticos (ALB), contenedores Docker (AWS ECS Fargate) y migración de la base de datos a AWS RDS PostgreSQL.")

    p_t3 = doc.add_paragraph()
    p_t3.add_run("3. Proyecto de Ciberseguridad y Cumplimiento Normativo (Fintech): ").font.bold = True
    p_t3.add_run("Implementación de autenticación de doble factor (MFA / SMS OTP), cifrado de extremo a extremo, blindaje contra ataques OWASP Top 10 con AWS WAF y políticas de privacidad conforme a la Ley de Protección de Datos Personales.")

    p_t4 = doc.add_paragraph()
    p_t4.add_run("4. Proyecto de Integración de Sistemas e Interoperabilidad (APIs): ").font.bold = True
    p_t4.add_run("Integración mediante microservicios y webhooks con burós de crédito (validación de riesgo) y pasarelas de desembolso interbancario (BCP, BBVA, Interbank, Yape/Plin).")

    # =========================================================================
    # SECCIÓN 3: 3 COMPONENTES CLAVE DE INFRAESTRUCTURA
    # =========================================================================
    h3 = doc.add_heading("3. Identificación y Justificación de 3 Componentes Clave de Infraestructura", level=1)
    h3.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    # Componente 1
    p_c1 = doc.add_paragraph()
    p_c1.add_run("A. Base de Datos Relacional Gestionada (DBaaS) - AWS RDS PostgreSQL Multi-AZ\n").font.bold = True
    p_c1.add_run(
        "• Rol: Almacenamiento seguro, consistente y transaccional (cumplimiento ACID) de clientes, créditos, cuotas y auditoría.\n"
        "• Justificación Técnica: Réplica síncrona automática en una segunda zona de disponibilidad con conmutación por error (Failover) en < 60s, "
        "backups continuos con Point-In-Time Recovery (PITR) de hasta 35 días y cifrado nativo KMS (AES-256)."
    )

    # Componente 2
    p_c2 = doc.add_paragraph()
    p_c2.add_run("B. Capa de Cómputo Elástico en Contenedores - AWS ECS con AWS Fargate\n").font.bold = True
    p_c2.add_run(
        "• Rol: Ejecución desacoplada de la API REST de créditos, motor de scoring y servicios backend.\n"
        "• Justificación Técnica: Escalado automático dinámico según la concurrencia de solicitudes, aislamiento por contenedores Docker inmutables "
        "y modelo Serverless sin necesidad de administrar servidores físicos."
    )

    # Componente 3
    p_c3 = doc.add_paragraph()
    p_c3.add_run("C. Red Privada Virtual Aislada, WAF y Balanceador - AWS VPC + AWS WAF + ALB\n").font.bold = True
    p_c3.add_run(
        "• Rol: Blindaje perimetral y aislamiento estricto de la base de datos en subredes privadas sin IP pública.\n"
        "• Justificación Técnica: Protección activa contra inyecciones SQL, ataques DDoS y XSS; gestión centralizada de certificados SSL/TLS 1.3 "
        "y distribución balanceada del tráfico HTTP/HTTPS hacia las instancias saludables."
    )

    doc.add_page_break()

    # =========================================================================
    # SECCIÓN 4: MATRIZ DE ROLES Y ASIGNACIÓN DE PRIMERA TAREA (RACI)
    # =========================================================================
    h4 = doc.add_heading("4. Matriz de Roles, Responsabilidades y Asignación de la Primera Tarea", level=1)
    h4.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    table_roles = doc.add_table(rows=6, cols=3)
    table_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_roles = ["Rol en el Proyecto", "Primera Tarea Asignada (Hito de Inicio)", "Entregable Concreto"]
    
    # Encabezados
    for col_idx, text in enumerate(headers_roles):
        cell = table_roles.rows[0].cells[col_idx]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, 100, 100, 120, 120)

    roles_data = [
        ("1. Project Manager (Scrum Master)", "Elaborar el Acta de Constitución del Proyecto (Project Charter) y el Backlog de Sprints inicial.", "Project Charter firmado y tablero Jira."),
        ("2. Arquitecto Cloud / DevOps", "Diseñar la topología de red y aprovisionar la Red Privada (VPC, subredes seguras e IAM).", "Scripts Terraform de infraestructura base."),
        ("3. Administrador de BD (DBA)", "Ejecutar el diagnóstico (Assessment) de la BD local: volumetría, tipos de datos y script DDL.", "Reporte de auditoría y script DDL depurado."),
        ("4. Desarrollador Full Stack", "Crear el repositorio base e implementar el cotizador interactivo con Amortización Francesa.", "Prototipo funcional de la interfaz en Git."),
        ("5. Especialista en Seguridad / QA", "Definir la matriz de requisitos de seguridad (KMS/TLS) y el plan maestro de casos de prueba.", "Documento de Políticas y Matriz QA.")
    ]

    for row_idx, r_data in enumerate(roles_data, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(r_data):
            cell = table_roles.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            if col_idx == 0:
                run.font.bold = True
            set_cell_background(cell, bg)
            set_cell_margins(cell, 80, 80, 100, 100)

    doc.add_paragraph("\n")

    # =========================================================================
    # SECCIÓN 5: CICLO DE VIDA CON 1 ENTREGABLE POR FASE
    # =========================================================================
    h5 = doc.add_heading("5. Esquema del Ciclo de Vida del Proyecto (1 Entregable Clave por Fase)", level=1)
    h5.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    table_fases = doc.add_table(rows=6, cols=3)
    table_fases.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_fases = ["Fase del Ciclo de Vida", "Objetivo Principal", "Entregable Único y Obligatorio"]

    for col_idx, text in enumerate(headers_fases):
        cell = table_fases.rows[0].cells[col_idx]
        run = cell.paragraphs[0].add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "0F766E")
        set_cell_margins(cell, 100, 100, 120, 120)

    fases_data = [
        ("Fase 1: Inicio y Requerimientos", "Definir alcance del negocio, reglas de crédito y viabilidad técnica.", "Documento SRS (Especificación de Requerimientos) y Project Charter."),
        ("Fase 2: Diseño y Arquitectura", "Diseñar topología Cloud, modelo de datos y wireframes UX.", "Plan Maestro de Arquitectura Cloud y Especificación de Migración de BD."),
        ("Fase 3: Construcción y Migración", "Desarrollar la app web y transferir datos locales a la nube.", "Plataforma Web Funcional y BD Migrada en Entorno Staging."),
        ("Fase 4: Pruebas y Seguridad (QA)", "Validar precisión de cuotas, pruebas de carga y ciberseguridad.", "Informe de Certificación QA y Auditoría de Seguridad (OWASP)."),
        ("Fase 5: Despliegue y Cierre", "Ventana de corte (Cutover), pase a vivo y transferencia operativa.", "Acta Oficial de Pase a Producción (Go-Live) y Manual de Operaciones.")
    ]

    for row_idx, f_data in enumerate(fases_data, start=1):
        bg = "F0FDF4" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(f_data):
            cell = table_fases.rows[row_idx].cells[col_idx]
            run = cell.paragraphs[0].add_run(text)
            if col_idx == 0:
                run.font.bold = True
            elif col_idx == 2:
                run.font.bold = True
                run.font.color.rgb = RGBColor(15, 118, 110)
            set_cell_background(cell, bg)
            set_cell_margins(cell, 80, 80, 100, 100)

    doc.add_page_break()

    # =========================================================================
    # SECCIÓN 6: MANUAL DE USO Y FUNCIONAMIENTO PASO A PASO
    # =========================================================================
    h6 = doc.add_heading("6. Manual de Uso y Funcionamiento de la Plataforma Web Móvil", level=1)
    h6.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "La plataforma web desarrollada (PrestaFast) opera bajo un flujo transaccional continuo de 4 pasos sin recargas de página, optimizado para celulares y computadoras:"
    )

    steps_manual = [
        ("Paso 1 - Cotización y Amortización Francesa:", 
         "El usuario selecciona el monto mediante slider o botones rápidos (chips: S/ 1,000, S/ 3,000, S/ 5,000, S/ 10,000, S/ 20,000) y el plazo (6 a 36 meses). "
         "El motor calcula en tiempo real la cuota fija mensual mediante la fórmula bancaria francesa:\n"
         "   C = P * [ r * (1 + r)^n ] / [ (1 + r)^n - 1 ]\n"
         "Dispone además del botón 'Ver Tabla de Cuotas Detallada' para auditar cuota por cuota el desglose de capital, intereses y saldo remanente."),
        
        ("Paso 2 - Formulario de Datos y Scoring Crediticio:",
         "El cliente ingresa su DNI, nombres, teléfono celular, correo e ingresos mensuales. Al presionar 'Evaluar Crédito en Tiempo Real', "
         "el sistema ejecuta la consulta simulada contra el buró de crédito vía API Cloud y verifica la capacidad de endeudamiento."),
        
        ("Paso 3 - Validación de Seguridad en 2 Pasos (MFA / SMS OTP):",
         "Para garantizar la autenticidad del solicitante, se envía un código dinámico de 4 dígitos (SMS OTP) con temporizador regresivo de 45 segundos. "
         "Las 4 casillas cuentan con auto-enfoque al escribir para máxima agilidad en celulares."),
        
        ("Paso 4 - Selección de Entidad de Desembolso y Firma Digital:",
         "El usuario elige el banco de destino (BCP, BBVA, Interbank o Yape/Plin), ingresa su número de cuenta/CCI y acepta el pagaré digital con firma electrónica cifrada. "
         "Al confirmar, la transacción se procesa y se escribe directamente en la base de datos AWS RDS."),
        
        ("Pantalla de Comprobante Oficial:",
         "Se genera un ticket bancario con N° de Operación único (ej. OP-948102), fecha y hora exacta, datos del titular y Hash Criptográfico en la Nube. "
         "Incluye botón para Descargar o Imprimir el Comprobante en PDF."),
        
        ("Consola de Monitoreo Cloud en Tiempo Real:",
         "En la sección inferior, la tabla de auditoría en vivo muestra inmediatamente la nueva transacción registrada en AWS RDS PostgreSQL junto con las métricas de estado (Online Multi-AZ y cifrado KMS).")
    ]

    for title, desc in steps_manual:
        p = doc.add_paragraph()
        r_t = p.add_run(f"• {title}\n")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(37, 99, 235)
        p.add_run(desc)

    doc.add_paragraph("\n")

    # =========================================================================
    # SECCIÓN 7: ESTRATEGIA DE MIGRACIÓN DE BD Y SCRIPTS
    # =========================================================================
    h7 = doc.add_heading("7. Plan Técnico de Migración de Base de Datos y Validación por Checksums", level=1)
    h7.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    doc.add_paragraph(
        "Se adoptó la estrategia Replatform (Lift, Tinker and Shift) automatizada mediante script Python (scripts/migrate_to_cloud.py). "
        "El procedimiento garantiza integridad total mediante comparación de Checksums criptográficos SHA-256:"
    )

    doc.add_paragraph(
        "1. Extracción y Diagnóstico: Lectura de clientes y préstamos en BD local On-Premise.\n"
        "2. Checksum de Origen: Generación de hash SHA-256 de las tablas locales.\n"
        "3. Aprovisionamiento Cloud: Creación de tablas e índices en AWS RDS PostgreSQL.\n"
        "4. Carga y Sincronización: Inserción de registros y log de auditoría.\n"
        "5. Validación de Integridad: Comparación del Checksum Local vs. Checksum Cloud. La migración se certifica únicamente con coincidencia del 100%."
    )

    # =========================================================================
    # SECCIÓN 8: CONCLUSIONES Y ENLACES OFICIALES
    # =========================================================================
    h8 = doc.add_heading("8. Enlaces de Acceso y Conclusiones", level=1)
    h8.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    p_links = doc.add_paragraph()
    p_links.add_run("• Aplicación Web en Vivo: ").font.bold = True
    p_links.add_run("https://gmph2007.github.io/prestamos-rapidos-cloud/\n")
    p_links.add_run("• Repositorio de Código Fuente y Documentación: ").font.bold = True
    p_links.add_run("https://github.com/GMPH2007/prestamos-rapidos-cloud")

    # Guardar documento
    output_path = "DOCUMENTO_FINAL_PRESTAMOS_CLOUD.docx"
    doc.save(output_path)
    print(f"Documento Word creado con éxito en: {output_path}")

if __name__ == "__main__":
    create_word_document()
